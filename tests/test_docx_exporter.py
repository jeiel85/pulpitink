from pathlib import Path

from docx import Document

from pulpit_ink.core.export import DocxExporter, ExportRequest
from pulpit_ink.core.transcription.base import TranscriptionResult, TranscriptSegment
from pulpit_ink.services.settings_service import Settings, SettingsService


def _sample_result(tmp_path: Path) -> TranscriptionResult:
    return TranscriptionResult(
        source_path=Path("sermon.mp3"),
        audio_path=tmp_path / "processed.wav",
        language="ko",
        model_name="small",
        preset="sermon",
        segments=[
            TranscriptSegment(start=0.0, end=2.5, text="창세기 1장 1절 태초에 하나님이 천지를 창조하시니라.", speaker="목사님"),
            TranscriptSegment(start=2.5, end=5.0, text="이것은 두 번째 말씀 구절입니다.", speaker="목사님"),
        ],
        duration=5.0,
    )


def _make_request(result, tmp_path: Path, bible_refs=None) -> ExportRequest:
    refs = bible_refs if bible_refs is not None else ["창세기 1장 1절"]
    return ExportRequest(
        result=result,
        output_dir=tmp_path / "out",
        base_name="sermon",
        bible_refs=refs
    )


def test_docx_exporter_pulpit_desk_style(tmp_path: Path):
    # Set preference to pulpit_desk
    settings_service = SettingsService()
    settings_service.save(Settings(docx_template_style="pulpit_desk"))

    result = _sample_result(tmp_path)
    req = _make_request(result, tmp_path)

    path = DocxExporter().export(req)
    assert path.exists()

    # Load and verify document
    doc = Document(str(path))

    # In pulpit_desk:
    # 1. Large title
    # 2. Large body text
    # 3. No timestamps or bold speakers in paragraph texts
    assert len(doc.paragraphs) > 0
    text_content = [p.text for p in doc.paragraphs if p.text.strip()]

    # Title must be there
    assert "Sermon" in text_content[0]
    # Bible highlight box text must be there
    assert "[ 성경 본문 ]" in text_content[1]
    assert "창세기 1장 1절" in text_content[1]
    # Main segments
    assert "태초에 하나님이 천지를 창조하시니라." in text_content[2]
    # Verify no speakers like "[목사님]" or timestamps in pulpit_desk style text
    assert "[목사님]" not in text_content[2]


def test_docx_exporter_church_bulletin_style(tmp_path: Path):
    # Set preference to church_bulletin
    settings_service = SettingsService()
    settings_service.save(Settings(docx_template_style="church_bulletin"))

    result = _sample_result(tmp_path)
    req = _make_request(result, tmp_path)

    path = DocxExporter().export(req)
    assert path.exists()

    doc = Document(str(path))

    # In church_bulletin:
    # Speakers should be bold and prefixed like "[목사님]" and timestamps like "(00:00)"
    has_speaker = False
    has_timestamp = False
    for p in doc.paragraphs:
        if "[목사님]" in p.text:
            has_speaker = True
        if "(00:00)" in p.text or "(00:02)" in p.text:
            has_timestamp = True

    assert has_speaker
    assert has_timestamp


def test_docx_exporter_grid_review_style(tmp_path: Path):
    # Set preference to grid_review
    settings_service = SettingsService()
    settings_service.save(Settings(docx_template_style="grid_review"))

    result = _sample_result(tmp_path)
    req = _make_request(result, tmp_path)

    path = DocxExporter().export(req)
    assert path.exists()

    doc = Document(str(path))

    # In grid_review:
    # There must be a table
    assert len(doc.tables) == 1
    table = doc.tables[0]

    # Verify header and rows
    assert len(table.rows) == 3  # Header + 2 segments
    assert len(table.columns) == 4

    hdr_cells = [cell.text for cell in table.rows[0].cells]
    assert "번호" in hdr_cells
    assert "시간" in hdr_cells
    assert "화자" in hdr_cells
    assert "변환 내용" in hdr_cells[3]

    row1_cells = [cell.text for cell in table.rows[1].cells]
    assert row1_cells[0] == "1"
    assert "00:00" in row1_cells[1]
    assert row1_cells[2] == "목사님"
    assert "창세기 1장 1절" in row1_cells[3]


def test_docx_exporter_scripture_highlight_fallback(tmp_path: Path):
    # Test fallback extraction when bible_refs is empty
    settings_service = SettingsService()
    settings_service.save(Settings(docx_template_style="pulpit_desk"))

    result = _sample_result(tmp_path)
    # Pass empty bible_refs list
    req = _make_request(result, tmp_path, bible_refs=[])

    path = DocxExporter().export(req)
    assert path.exists()

    doc = Document(str(path))
    text_content = [p.text for p in doc.paragraphs if p.text.strip()]

    # Fallback should extract "창세기 1장 1절" from the text
    assert "[ 성경 본문 ]" in text_content[1]
    assert "창세기 1장 1절" in text_content[1]
