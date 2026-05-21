"""MS Word (.docx) Exporter implementation.

Uses python-docx to render professional sermon manuscripts with 3 custom templates
and a scripture highlight box on top.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from pulpit_ink.core.export.base import Exporter, ExportFormat, ExportRequest
from pulpit_ink.core.transcription.base import segment_display_text

logger = logging.getLogger("pulpit_ink.export.docx")


def set_paragraph_background(paragraph, color_hex: str):
    """Set the background color of a paragraph using openxml."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def set_paragraph_left_border(paragraph, color_hex: str, size_pt: float):
    """Set a thick left border on a paragraph to create a blockquote effect."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(int(size_pt * 8)))  # size in 1/8 pt
    left.set(qn('w:space'), '12')  # padding space in pt
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def format_duration(seconds: float) -> str:
    """Format seconds into HH:MM:SS or MM:SS."""
    total_sec = int(seconds)
    hours = total_sec // 3600
    minutes = (total_sec % 3600) // 60
    secs = total_sec % 60
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def set_cell_background(cell, color_hex: str):
    """Set table cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


class DocxExporter(Exporter):
    """Generates MS Word manuscripts from transcription results."""

    format = ExportFormat.DOCX

    def export(self, request: ExportRequest) -> Path:
        request.output_dir.mkdir(parents=True, exist_ok=True)
        target = request.output_dir / f"{request.base_name}.docx"

        # Load user preference for the template style
        from pulpit_ink.services.settings_service import SettingsService
        settings = SettingsService().load()
        style = settings.docx_template_style  # "pulpit_desk", "church_bulletin", "grid_review"
        logger.info("Exporting Word document using style: %s", style)

        doc = Document()

        # 1. Set general document parameters based on style
        sections = doc.sections
        for section in sections:
            if style == "pulpit_desk":
                # Margin: Top/Bottom 30mm (1.18"), Left/Right 25mm (0.98")
                section.top_margin = Inches(1.18)
                section.bottom_margin = Inches(1.18)
                section.left_margin = Inches(0.98)
                section.right_margin = Inches(0.98)
            else:
                # Standard margin: 1 inch everywhere
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

        # 2. Add Document Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(24)
        run = title_p.add_run(request.base_name.replace("_", " ").title())
        run.bold = True
        run.size = Pt(22) if style == "pulpit_desk" else Pt(18)
        run.font.name = "맑은 고딕" if style != "pulpit_desk" else "바탕"

        # 3. Add Scripture Highlight Box if bible references exist
        bible_refs = list(request.bible_refs)
        if not bible_refs:
            # Fallback: Extract from full text if empty
            from pulpit_ink.core.reference.parser import _extract_bible_refs
            bible_refs = _extract_bible_refs(request.result.full_text)

        if bible_refs:
            # Create the blockquote-style container paragraph
            box_p = doc.add_paragraph()
            box_p.paragraph_format.left_indent = Inches(0.5)
            box_p.paragraph_format.right_indent = Inches(0.5)
            box_p.paragraph_format.space_after = Pt(18)
            box_p.paragraph_format.space_before = Pt(12)

            # Apply styling (Light Beige background `#F8F6F0`, Navy border `#1E293B`)
            set_paragraph_background(box_p, "F8F6F0")
            set_paragraph_left_border(box_p, "1E293B", size_pt=3.0)

            # Write content inside the box
            icon_run = box_p.add_run("📖 [ 성경 본문 ]\n\n")
            icon_run.bold = True
            icon_run.size = Pt(11)
            icon_run.font.color.rgb = RGBColor(30, 41, 59)
            icon_run.font.name = "맑은 고딕" if style != "pulpit_desk" else "바탕"

            for i, ref in enumerate(bible_refs):
                ref_run = box_p.add_run(f"\"{ref}\"" + ("\n" if i < len(bible_refs) - 1 else ""))
                ref_run.italic = True
                ref_run.size = Pt(10.5)
                ref_run.font.color.rgb = RGBColor(51, 65, 85)
                ref_run.font.name = "맑은 고딕" if style != "pulpit_desk" else "바탕"

        # 4. Render Segments according to Template Style
        if style == "grid_review":
            # Table-style Grid Review
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # Set styling
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_titles = ["번호", "시간", "화자", "변환 내용 (edited_text)"]
            # Column widths: 0.5", 1.2", 1.2", 4.1" (Total 7.0")
            widths = [Inches(0.5), Inches(1.2), Inches(1.2), Inches(4.1)]

            for idx, cell in enumerate(hdr_cells):
                cell.width = widths[idx]
                set_cell_background(cell, "1E3C72")  # Navy Blue header
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_run = p.add_run(hdr_titles[idx])
                p_run.bold = True
                p_run.font.color.rgb = RGBColor(255, 255, 255)
                p_run.font.name = "맑은 고딕"
                p_run.font.size = Pt(10)

            for i, seg in enumerate(request.result.segments, 1):
                row_cells = table.add_row().cells
                for idx, cell in enumerate(row_cells):
                    cell.width = widths[idx]

                # 번호
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                num_run = row_cells[0].paragraphs[0].add_run(str(i))
                num_run.font.name = "맑은 고딕"
                num_run.font.size = Pt(9.5)

                # 시간
                time_str = f"{format_duration(seg.start)} - {format_duration(seg.end)}"
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                time_run = row_cells[1].paragraphs[0].add_run(time_str)
                time_run.font.name = "맑은 고딕"
                time_run.font.size = Pt(9.5)

                # 화자
                speaker_str = seg.speaker or f"화자 {i}"
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                spk_run = row_cells[2].paragraphs[0].add_run(speaker_str)
                spk_run.bold = True
                spk_run.font.name = "맑은 고딕"
                spk_run.font.size = Pt(9.5)

                # 내용
                text_str = segment_display_text(seg)
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                txt_run = row_cells[3].paragraphs[0].add_run(text_str)
                txt_run.font.name = "맑은 고딕"
                txt_run.font.size = Pt(10)

        elif style == "church_bulletin":
            # Handout-style with small gray timestamps and bold speakers
            last_speaker = None
            for seg in request.result.segments:
                text_str = segment_display_text(seg)
                if not text_str.strip():
                    continue

                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.3
                p.paragraph_format.space_after = Pt(6)

                # 1. Speaker Tag (if changed)
                current_speaker = seg.speaker or "화자"
                if current_speaker != last_speaker:
                    spk_run = p.add_run(f"[{current_speaker}] ")
                    spk_run.bold = True
                    spk_run.font.name = "맑은 고딕"
                    spk_run.font.size = Pt(10.5)
                    last_speaker = current_speaker

                # 2. Small Gray Timestamp
                time_str = f"({format_duration(seg.start)}) "
                time_run = p.add_run(time_str)
                time_run.font.size = Pt(8.5)
                time_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
                time_run.font.name = "맑은 고딕"

                # 3. Transcript text
                txt_run = p.add_run(text_str)
                txt_run.font.name = "맑은 고딕"
                txt_run.font.size = Pt(10.5)

        else:
            # pulpit_desk (Pulpit Large Text) style
            # Clean layout, large serif text, line spacing 1.8, no timestamps/speakers
            for seg in request.result.segments:
                text_str = segment_display_text(seg)
                if not text_str.strip():
                    continue

                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.8
                p.paragraph_format.space_after = Pt(14)

                txt_run = p.add_run(text_str)
                txt_run.font.name = "바탕"
                txt_run.font.size = Pt(14)  # Large font size
                txt_run.font.color.rgb = RGBColor(15, 23, 42)  # Dark slate gray

        doc.save(str(target))
        return target
